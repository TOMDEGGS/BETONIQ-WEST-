from docx import Document
from docx.shared import RGBColor, Pt

PATH = "AVM_MORGAN_SECURITY_KEYNOTE_ADDRESS_V4_MOD_AUG31.docx"
doc = Document(PATH)

# Exact heading text -> slide label
LABELS = {
    "OPENING: ONE ECOSYSTEM, NOT SEPARATE PROBLEMS": "SLIDES 2\u20133",
    "SECTION ONE: THE NORTH-CENTRAL SECURITY QUESTION": "SLIDE 4",
    "SECTION TWO: THE CENTRAL PROBLEM \u2014 WE ARE TOO REACTIVE": "SLIDE 5",
    "SECTION THREE: SECURITY AS COLLECTIVE RESPONSIBILITY": "SLIDE 6",
    "SECTION FOUR: ONE SECURITY PICTURE, MULTIPLE CAPABILITIES": "SLIDE 7",
    "SECTION FIVE: A REGIONAL SECURITY INTELLIGENCE AND FUSION CENTRE": "SLIDE 8",
    "SECTION SIX: FROM INTELLIGENCE COLLECTION TO INTELLIGENCE ACTION": "SLIDE 9",
    "SECTION SEVEN: THE BRITAIN AND ISRAEL MODEL \u2014 PROACTIVE INTELLIGENCE TO NIP THREATS IN THE BUD": "SLIDE 10",
    "SECTION EIGHT: THE COMMUNITY MUST BECOME THE FIRST SENSOR": "SLIDE 11",
    "SECTION NINE: TRADITIONAL RULERS IN THE SECURITY ARCHITECTURE": "SLIDE 12",
    "SECTION TEN: SECURITY CORRIDORS, NOT ADMINISTRATIVE BOUNDARIES": "SLIDE 13",
    "SECTION ELEVEN: RAPID RESPONSE \u2014 THE THIRTY TO SIXTY MINUTE WINDOW": "SLIDE 14",
    "SECTION TWELVE: TECHNOLOGY SHOULD MULTIPLY HUMAN CAPABILITY": "SLIDE 15",
    "SECTION THIRTEEN: THE FOREST PROBLEM": "SLIDE 16",
    "SECTION FOURTEEN: SEPARATE THE CRIMINAL FROM THE COMMUNITY": "SLIDE 17",
    "SECTION FIFTEEN: FARMER-HERDER CONFLICT AS A GOVERNANCE ISSUE": "SLIDE 18",
    "SECTION SIXTEEN: YOUNG PEOPLE \u2014 THE MOST IMPORTANT LONG-TERM INVESTMENT": "SLIDE 19",
    "SECTIONS SEVENTEEN AND EIGHTEEN: WOMEN AS SECURITY PARTNERS, AND THE ROLE OF LOCAL GOVERNMENT": "SLIDE 20",
    "SECTION NINETEEN: STATE POLICE \u2014 AN OPPORTUNITY THAT REQUIRES SAFEGUARDS": "SLIDE 21",
    "SECTIONS TWENTY AND TWENTY-ONE: PREDICTABLE FUNDING AND PERSONNEL WELFARE": "SLIDE 22",
    "SECTION TWENTY-TWO: THE THREE-LAYER SECURITY MODEL": "SLIDE 23",
    "SECTION TWENTY-THREE: A TWENTY-FOUR-HOUR SECURITY CYCLE": "SLIDE 24",
    "SECTION TWENTY-FOUR: WHAT SHOULD HAPPEN IN THE NEXT NINETY DAYS": "SLIDE 25",
    "SECTION TWENTY-FIVE: THE NORTH-CENTRAL SECURITY COMPACT \u2014 SEVEN COMMITMENTS": "SLIDE 26",
    "SECTION TWENTY-SIX: WHAT SUCCESS SHOULD LOOK LIKE IN THREE YEARS": "SLIDE 27",
    "SECTION TWENTY-SEVEN: THE BIGGEST CHANGE WE NEED IS MENTAL": "SLIDE 28",
    "CLOSING": "SLIDE 29",
    "REFERENCES": "SLIDE 30",
}

BLUE = RGBColor(0x1F, 0x4E, 0x9C)
matched = []
missed = list(LABELS.keys())

for p in doc.paragraphs:
    txt = p.text.strip()
    if txt in LABELS:
        label = LABELS[txt]
        run = p.add_run(f"   [{label}]")
        run.bold = True
        run.font.color.rgb = BLUE
        # Match size roughly to the heading's own run size if available, else default 14pt
        try:
            base_size = p.runs[0].font.size
            run.font.size = base_size if base_size else Pt(14)
        except Exception:
            run.font.size = Pt(14)
        matched.append(txt)
        if txt in missed:
            missed.remove(txt)

# Add SLIDE 1 marker under the title block (after the subtitle line, before the greeting)
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Towards a People-Centred, Intelligence-Led and Sustainable Security Architecture":
        run = p.add_run("   [SLIDE 1]")
        run.bold = True
        run.font.color.rgb = BLUE
        run.font.size = Pt(12)
        matched.append("TITLE_MARKER")
        break

doc.save(PATH)
print(f"Labeled {len(matched)} headings.")
print("Missed:", missed)
